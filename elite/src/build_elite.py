# -*- coding: utf-8 -*-
"""Sinh 4 tệp Word cho hai bộ đề VẬN DỤNG CAO – VẬT LÍ 12.

  BỘ 1  (Chương I + Chương II): 1 tệp đề bài + 1 tệp lời giải
  BỘ 2  (Chương III + Chương IV): 1 tệp đề bài + 1 tệp lời giải

Mã số câu hỏi được giữ nguyên tuyệt đối giữa tệp đề bài và tệp lời giải.
"""
import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
FIGS = os.path.join(ROOT, "figs")
sys.path.insert(0, HERE)

import book1
import book2

FONT = "Times New Roman"
LETTERS = ["A", "B", "C", "D"]
DS_LABELS = ["a)", "b)", "c)", "d)"]

DARKRED = RGBColor(0xA6, 0x1B, 0x1B)
DARKBLUE = RGBColor(0x1F, 0x3A, 0x6E)
DARKGREEN = RGBColor(0x1E, 0x64, 0x36)
GREYTXT = RGBColor(0x55, 0x55, 0x55)


# ------------------------------------------------------------------ tiện ích
def set_base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(3)
    st.paragraph_format.line_spacing = 1.13
    for s in doc.sections:
        s.top_margin = Cm(1.9)
        s.bottom_margin = Cm(1.9)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(1.8)


def P(doc, text="", bold=False, italic=False, size=12, align=None, before=0, after=3,
      indent=None, color=None, first_line=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent is not None:
        pf.left_indent = Cm(indent)
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    if text:
        for i, line in enumerate(str(text).split("\n")):
            if i:
                p.add_run().add_break()
            r = p.add_run(line)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
            r.font.name = FONT
            if color is not None:
                r.font.color.rgb = color
    return p


def rich(doc, parts, indent=None, before=0, after=3, size=12, align=None):
    """parts = [(text, bold, italic, color|None), ...]"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent is not None:
        pf.left_indent = Cm(indent)
    for text, b, i, c in parts:
        for k, line in enumerate(str(text).split("\n")):
            if k:
                p.add_run().add_break()
            r = p.add_run(line)
            r.bold = b
            r.italic = i
            r.font.size = Pt(size)
            r.font.name = FONT
            if c is not None:
                r.font.color.rgb = c
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def box(doc, title, text, fill="EEF3FA", tcolor=DARKBLUE, size=11.5):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, fill)
    c.paragraphs[0].paragraph_format.space_after = Pt(2)
    if title:
        r = c.paragraphs[0].add_run(title)
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = FONT
        r.font.color.rgb = tcolor
    tp = c.add_paragraph()
    tp.paragraph_format.space_after = Pt(1)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            tp.add_run().add_break()
        r = tp.add_run(line)
        r.font.size = Pt(size)
        r.font.name = FONT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def picture(doc, name, caption=None, maxw=12.5, maxh=8.2):
    path = os.path.join(FIGS, name + ".png")
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    with Image.open(path) as im:
        w, h = im.size
    width = maxw
    if h / w * width > maxh:
        width = maxh * w / h
    doc.add_picture(path, width=Cm(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(4)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
    if caption:
        P(doc, caption, italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
          after=6, color=GREYTXT)


def table(doc, caption, headers, rows, size=10.5):
    if caption:
        P(doc, caption, italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
          before=4, after=2, color=GREYTXT)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        c = t.rows[0].cells[j]
        shade(c, "DCE6F1")
        c.paragraphs[0].paragraph_format.space_after = Pt(1)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(htxt)
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = FONT
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].paragraph_format.space_after = Pt(1)
            if j:
                cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cells[j].paragraphs[0].add_run(str(val))
            r.font.size = Pt(size)
            r.font.name = FONT
            if j == 0:
                r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_page_numbers(doc):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.size = Pt(9)
        run.font.name = FONT
        for instr in ("begin", "PAGE", "end"):
            el = OxmlElement("w:fldChar") if instr in ("begin", "end") \
                else OxmlElement("w:instrText")
            if instr in ("begin", "end"):
                el.set(qn("w:fldCharType"), instr)
            else:
                el.set(qn("xml:space"), "preserve")
                el.text = " PAGE "
            run._r.append(el)


# ------------------------------------------------------------ cân bằng đáp án
def _numeric_options(opts):
    n = 0
    for o in opts:
        t = o.strip().lstrip("≈±").lstrip()
        if t[:1].isdigit() or t[:1] in "-−":
            n += 1
    return n >= 3


def rebalance(items):
    """Hoán vị vị trí phương án đúng để đáp án phân bố đều giữa A, B, C, D.

    Chỉ đụng tới những câu có phương án dạng chữ; các câu có phương án là giá trị
    số được giữ nguyên thứ tự tăng dần cho dễ đọc. Nội dung phương án không đổi,
    chỉ đổi chỗ, nên lời giải (vốn không nhắc tới chữ cái) vẫn đúng nguyên vẹn.
    """
    numeric, pool = [], []
    for it in items:
        (numeric if _numeric_options(it["o"]) else pool).append(it)

    # Với bộ phương án là giá trị số, chỉ cho phép ĐẢO CHIỀU sắp xếp (tăng dần ↔
    # giảm dần) để vẫn dễ đọc; phép đảo này đưa đáp án A↔D và B↔C.
    fixed = {L: 0 for L in LETTERS}
    for it in numeric:
        fixed[it["a"]] += 1
    for _pass in range(4):
        changed = False
        for it in numeric:
            cur = it["a"]
            mir = LETTERS[3 - LETTERS.index(cur)]
            if fixed[cur] > fixed[mir] + 1:
                it["o"].reverse()
                it["a"] = mir
                fixed[cur] -= 1
                fixed[mir] += 1
                changed = True
        if not changed:
            break

    alloc = {L: 0 for L in LETTERS}
    for _ in pool:
        L = min(LETTERS, key=lambda x: (fixed[x] + alloc[x], LETTERS.index(x)))
        alloc[L] += 1
    slots = []
    for L in LETTERS:
        slots += [L] * alloc[L]

    for it in pool:                       # ưu tiên giữ nguyên nếu còn chỗ
        if it["a"] in slots:
            slots.remove(it["a"])
        else:
            it["_want"] = None
    for it in pool:
        if it.get("_want", "keep") is None:
            it["_want"] = slots.pop(0)
            i, j = LETTERS.index(it["a"]), LETTERS.index(it["_want"])
            it["o"][i], it["o"][j] = it["o"][j], it["o"][i]
            it["a"] = it["_want"]
        it.pop("_want", None)

    out = {L: 0 for L in LETTERS}
    for it in items:
        out[it["a"]] += 1
    return out


# ---------------------------------------------------------------- kết cấu đề
def kind_of(item):
    if "o" in item:
        return "mc"
    if "items" in item:
        return "ds"
    return "sa"


PART_TITLES = [
    ("I", "PHẦN I – CÂU TRẮC NGHIỆM NHIỀU PHƯƠNG ÁN LỰA CHỌN",
     "Mỗi câu hỏi có bốn phương án lựa chọn; thí sinh chỉ chọn một phương án."),
    ("II", "PHẦN II – CÂU TRẮC NGHIỆM ĐÚNG/SAI",
     "Mỗi câu có bốn ý a), b), c), d); ở mỗi ý thí sinh chọn ĐÚNG hoặc SAI."),
    ("III", "PHẦN III – CÂU TRẮC NGHIỆM TRẢ LỜI NGẮN",
     "Thí sinh ghi đáp số vào ô trả lời, làm tròn theo đúng yêu cầu của đề."),
]


def structure(book):
    """[(mã phần, tiêu đề, mô tả, [(mã câu, item)])]"""
    out = []
    for (code, title, note), items in zip(PART_TITLES, (book.P1, book.P2, book.P3)):
        lst = [("Câu %s.%d" % (code, i), it) for i, it in enumerate(items, 1)]
        out.append((code, title, note, lst))
    return out


def short_answer(item):
    k = kind_of(item)
    if k == "mc":
        return item["a"]
    if k == "ds":
        return "".join("Đ" if v else "S" for _t, v, _e in item["items"])
    return item["ans"]


# ------------------------------------------------------------------- hiển thị
def render_tag(doc, item):
    if item.get("tag"):
        rich(doc, [("◆ " + item["tag"], False, True, DARKGREEN)],
             indent=0.0, before=8, after=1, size=10)


def render_question(doc, ident, item):
    k = kind_of(item)
    render_tag(doc, item)
    stem = item["stem"] if k == "ds" else item["q"]
    rich(doc, [(ident + ". ", True, False, DARKRED), (stem, False, False, None)],
         before=0, after=2)
    if item.get("tbl"):
        table(doc, *item["tbl"])
    if item.get("fig"):
        picture(doc, item["fig"], item.get("cap"))
    if k == "ds":
        for lab, (txt, _v, _e) in zip(DS_LABELS, item["items"]):
            rich(doc, [(lab + " ", True, False, None), (txt, False, False, None)],
                 indent=0.6, after=1)
    elif k == "mc":
        for i, opt in enumerate(item["o"]):
            rich(doc, [("%s. " % LETTERS[i], True, False, None), (opt, False, False, None)],
                 indent=0.6, after=0)
    else:
        P(doc, "Trả lời:  ......................................................",
          size=11.5, indent=0.6, after=1)


def render_solution(doc, ident, item):
    k = kind_of(item)
    render_tag(doc, item)
    stem = item["stem"] if k == "ds" else item["q"]
    rich(doc, [(ident + ". ", True, False, DARKRED), (stem, False, True, None)],
         before=0, after=2)
    if item.get("tbl"):
        table(doc, *item["tbl"])
    if item.get("fig"):
        picture(doc, item["fig"], item.get("cap"), maxw=11.0, maxh=7.2)

    if k == "ds":
        for lab, (txt, val, ex) in zip(DS_LABELS, item["items"]):
            rich(doc, [(lab + " ", True, False, None), (txt, False, False, None),
                       ("   →  " + ("ĐÚNG" if val else "SAI"), True, False,
                        DARKBLUE if val else DARKRED)], indent=0.6, after=1)
            rich(doc, [(ex, False, False, None)], indent=1.1, after=3, size=11.5)
        return

    if k == "mc":
        ans = item["a"]
        txt = item["o"][LETTERS.index(ans)]
        rich(doc, [("Đáp án: %s. " % ans, True, False, DARKRED), (txt, True, False, DARKRED)],
             indent=0.6, after=2)
    else:
        rich(doc, [("Đáp án: ", True, False, DARKRED), (item["ans"], True, False, DARKRED)],
             indent=0.6, after=2)
    rich(doc, [("Lời giải chi tiết:", True, True, DARKBLUE)], indent=0.6, after=1, size=11.5)
    rich(doc, [(item["sol"], False, False, None)], indent=0.6, after=3, size=11.5)


def answer_key(doc, title, pairs, ncol=6):
    P(doc, title, bold=True, size=12, before=8, after=4, color=DARKBLUE)
    nrow = (len(pairs) + ncol - 1) // ncol
    t = doc.add_table(rows=nrow, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (ident, ans) in enumerate(pairs):
        c = t.rows[i // ncol].cells[i % ncol]
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(ident.replace("Câu ", ""))
        r.font.size = Pt(9.5)
        r.font.name = FONT
        c.paragraphs[0].add_run(": ").font.size = Pt(9.5)
        r2 = c.paragraphs[0].add_run(ans)
        r2.bold = True
        r2.font.size = Pt(9.5)
        r2.font.name = FONT
        r2.font.color.rgb = DARKRED
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def skill_matrix(doc, struct):
    P(doc, "BẢNG ĐỊNH HƯỚNG KĨ NĂNG CỦA TỪNG CÂU", bold=True, size=12.5,
      align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4, color=DARKBLUE)
    rows = []
    for _code, _title, _note, lst in struct:
        for ident, it in lst:
            rows.append([ident.replace("Câu ", ""), it.get("tag", "")])
    half = (len(rows) + 1) // 2
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(["Câu", "Kĩ năng trọng tâm", "Câu", "Kĩ năng trọng tâm"]):
        c = t.rows[0].cells[j]
        shade(c, "DCE6F1")
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraphs[0].paragraph_format.space_after = Pt(1)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = FONT
    for i in range(half):
        cells = t.add_row().cells
        pair = [rows[i]] + ([rows[i + half]] if i + half < len(rows) else [["", ""]])
        vals = pair[0] + pair[1]
        for j, v in enumerate(vals):
            cells[j].paragraphs[0].paragraph_format.space_after = Pt(0)
            if j % 2 == 0:
                cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cells[j].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
            r.font.name = FONT
            if j % 2 == 0:
                r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


ORIENT = (
    "Tài liệu bám sát Chương trình GDPT 2018 và cấu trúc định dạng đề thi tốt nghiệp THPT áp dụng từ năm "
    "2025, đồng thời hướng tới các kì thi đánh giá năng lực. Toàn bộ câu hỏi được biên soạn MỚI, ở mức "
    "VẬN DỤNG và VẬN DỤNG CAO.\n"
    "Nguyên tắc biên soạn: câu hỏi khó vì tư duy vật lí, không khó vì tính toán rườm rà. Mọi kiến thức và "
    "công cụ toán học đều nằm trong chương trình lớp 12; không dùng giải tích, phương trình vi phân hay "
    "kiến thức đại học.\n"
    "Trọng tâm đánh giá: nhận ra điều kiện áp dụng của từng định luật, phát hiện điều kiện ẩn, đọc và biến "
    "đổi đồ thị, xử lí số liệu thực nghiệm, suy luận ngược, so sánh nhiều tình huống và tránh các bẫy khái "
    "niệm tinh vi."
)

HOWTO_Q = (
    "• Mỗi câu được gắn một nhãn kĩ năng (dòng in nghiêng màu xanh lá) cho biết năng lực trọng tâm mà câu "
    "đó kiểm tra; hãy tự đối chiếu sau khi làm bài.\n"
    "• Ở Phần I, nhiều câu có tới ba phương án nhiễu được thiết kế từ ba sai lầm khác nhau; chọn được đáp "
    "án đúng chưa đủ, hãy giải thích được vì sao ba phương án còn lại sai.\n"
    "• Ở Phần II, mỗi ý được chấm độc lập; đừng suy đoán theo kiểu “chắc phải có hai ý đúng, hai ý sai”.\n"
    "• Ở Phần III, hãy chú ý đơn vị và yêu cầu làm tròn ghi ở cuối mỗi câu."
)

HOWTO_S = (
    "• Mã số câu trong tệp này trùng khớp tuyệt đối với tệp đề bài tương ứng.\n"
    "• Với câu đúng/sai, đáp án được ghi theo thứ tự bốn ý a) b) c) d), ví dụ “ĐĐSĐ”.\n"
    "• Mỗi lời giải đều nêu rõ mấu chốt vật lí, các bước tính và — khi cần — phân tích bẫy khái niệm cùng "
    "lí do các cách làm khác cho kết quả sai.\n"
    "• Bảng đáp án nhanh của từng phần được đặt ngay đầu phần tương ứng."
)


def cover(doc, book_no, chapters, subtitle, kind_label):
    P(doc, "TÀI LIỆU BỒI DƯỠNG HỌC SINH GIỎI VÀ LUYỆN THI", bold=True, size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=0, color=GREYTXT)
    P(doc, "VẬT LÍ 12", bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, after=0,
      color=DARKRED)
    P(doc, "TUYỂN TẬP BÀI TẬP VẬN DỤNG CAO", bold=True, size=15,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=2, color=DARKRED)
    P(doc, "Định hướng kì thi tốt nghiệp THPT và đánh giá năng lực 2026 – 2027",
      italic=True, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    P(doc, "BỘ %d — %s" % (book_no, chapters), bold=True, size=16,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=2, color=DARKBLUE)
    P(doc, kind_label, bold=True, size=13.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    P(doc, subtitle, italic=True, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    box(doc, "ĐỊNH HƯỚNG BIÊN SOẠN", ORIENT, fill="F4F6F8")


# ------------------------------------------------------------------ xây tệp
def build_questions(book, book_no, chapters, filename, subtitle):
    struct = structure(book)
    doc = Document()
    set_base(doc)
    add_page_numbers(doc)
    cover(doc, book_no, chapters, subtitle, "TỆP ĐỀ BÀI")
    box(doc, "HƯỚNG DẪN SỬ DỤNG", HOWTO_Q, fill="F2F8F0", tcolor=DARKGREEN)
    skill_matrix(doc, struct)
    doc.add_page_break()

    total = 0
    for idx, (_code, title, note, lst) in enumerate(struct):
        P(doc, title, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
          before=6, after=2, color=DARKRED)
        P(doc, "(%d câu)  %s" % (len(lst), note), italic=True, size=10.5,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=6, color=GREYTXT)
        for ident, item in lst:
            render_question(doc, ident, item)
            total += 1
        if idx < len(struct) - 1:
            doc.add_page_break()
    P(doc, "--- HẾT ĐỀ BÀI ---", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=10)
    path = os.path.join(ROOT, filename)
    doc.save(path)
    return path, total


def build_solutions(book, book_no, chapters, filename, subtitle):
    struct = structure(book)
    doc = Document()
    set_base(doc)
    add_page_numbers(doc)
    cover(doc, book_no, chapters, subtitle, "TỆP ĐÁP ÁN VÀ LỜI GIẢI CHI TIẾT")
    box(doc, "CÁCH ĐỌC TÀI LIỆU NÀY", HOWTO_S, fill="F4F6F8")

    P(doc, "BẢNG ĐÁP ÁN TOÀN BỘ", bold=True, size=13,
      align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=4, color=DARKRED)
    for _code, title, _note, lst in struct:
        answer_key(doc, title, [(i, short_answer(it)) for i, it in lst],
                   ncol=6 if len(lst) > 10 else 5)
    doc.add_page_break()

    for idx, (_code, title, note, lst) in enumerate(struct):
        P(doc, "LỜI GIẢI – " + title, bold=True, size=14,
          align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=6, color=DARKRED)
        for ident, item in lst:
            render_solution(doc, ident, item)
        if idx < len(struct) - 1:
            doc.add_page_break()
    P(doc, "--- HẾT LỜI GIẢI ---", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=10)
    path = os.path.join(ROOT, filename)
    doc.save(path)
    return path


CH1 = "CHƯƠNG I: VẬT LÍ NHIỆT  •  CHƯƠNG II: KHÍ LÍ TƯỞNG"
CH2 = "CHƯƠNG III: TỪ TRƯỜNG  •  CHƯƠNG IV: VẬT LÍ HẠT NHÂN"
SUB1 = ("Cấu trúc chất – nội năng – định luật I nhiệt động lực học – nhiệt dung riêng, "
        "nhiệt nóng chảy, nhiệt hoá hơi – mô hình động học phân tử – các định luật chất khí")
SUB2 = ("Từ trường và lực từ – từ thông – cảm ứng điện từ và định luật Lenz – dòng điện xoay chiều và "
        "máy biến áp – cấu tạo hạt nhân, năng lượng liên kết – phóng xạ, phân hạch và nhiệt hạch")


if __name__ == "__main__":
    print("Phân bố đáp án Phần I sau khi cân bằng:")
    print("  BỘ 1:", rebalance(book1.P1))
    print("  BỘ 2:", rebalance(book2.P1))

    p, n = build_questions(book1, 1, CH1,
                           "BO1_DE_BAI_VAN_DUNG_CAO_CHUONG_1_VA_2.docx", SUB1)
    print("Đã tạo:", os.path.basename(p), "-", n, "câu")
    p = build_solutions(book1, 1, CH1,
                        "BO1_LOI_GIAI_CHI_TIET_CHUONG_1_VA_2.docx", SUB1)
    print("Đã tạo:", os.path.basename(p))

    p, n = build_questions(book2, 2, CH2,
                           "BO2_DE_BAI_VAN_DUNG_CAO_CHUONG_3_VA_4.docx", SUB2)
    print("Đã tạo:", os.path.basename(p), "-", n, "câu")
    p = build_solutions(book2, 2, CH2,
                        "BO2_LOI_GIAI_CHI_TIET_CHUONG_3_VA_4.docx", SUB2)
    print("Đã tạo:", os.path.basename(p))
