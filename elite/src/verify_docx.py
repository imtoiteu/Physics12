# -*- coding: utf-8 -*-
"""Kiểm tra lại các tệp .docx ĐÃ SINH RA: đồng bộ mã câu, đáp án, hình, bảng."""
import os
import re
import sys
import zipfile

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
sys.path.insert(0, HERE)

PAIRS = [
    ("BO1_DE_BAI_VAN_DUNG_CAO_CHUONG_1_VA_2.docx",
     "BO1_LOI_GIAI_CHI_TIET_CHUONG_1_VA_2.docx", "BỘ 1"),
    ("BO2_DE_BAI_VAN_DUNG_CAO_CHUONG_3_VA_4.docx",
     "BO2_LOI_GIAI_CHI_TIET_CHUONG_3_VA_4.docx", "BỘ 2"),
]

RE_ID = re.compile(r"^(Câu (?:I{1,3})\.\d+)\.")
RE_ANS = re.compile(r"^Đáp án:\s*(.+?)\s*$")
BAD = []


def paras(path):
    d = Document(path)
    return d, [p.text.strip() for p in d.paragraphs]


def ids_in(texts):
    out = []
    for t in texts:
        m = RE_ID.match(t)
        if m:
            out.append(m.group(1))
    return out


def images_in(path):
    with zipfile.ZipFile(path) as z:
        return [n for n in z.namelist() if n.startswith("word/media/")]


def main():
    import book1, book2
    import build_elite as B

    # tái tạo đúng trạng thái dữ liệu sau khi cân bằng đáp án
    B.rebalance(book1.P1)
    B.rebalance(book2.P1)
    books = {"BỘ 1": book1, "BỘ 2": book2}

    for qf, sf, label in PAIRS:
        qp, sp = os.path.join(ROOT, qf), os.path.join(ROOT, sf)
        for p in (qp, sp):
            if not os.path.exists(p):
                BAD.append("Thiếu tệp %s" % p)
                return
        dq, tq = paras(qp)
        ds, ts = paras(sp)
        qids, sids = ids_in(tq), ids_in(ts)

        book = books[label]
        struct = B.structure(book)
        expect = [ident for _c, _t, _n, lst in struct for ident, _it in lst]

        if qids != expect:
            BAD.append("%s: mã câu trong tệp đề bài không khớp danh sách chuẩn" % label)
        if sids != expect:
            BAD.append("%s: mã câu trong tệp lời giải không khớp danh sách chuẩn" % label)
        if qids != sids:
            BAD.append("%s: mã câu giữa đề bài và lời giải KHÔNG đồng bộ" % label)

        # đáp án trong tệp lời giải phải khớp dữ liệu nguồn
        got = []
        for t in ts:
            m = RE_ANS.match(t)
            if m:
                got.append(m.group(1))
        exp_ans = []
        for _c, _t, _n, lst in struct:
            for _ident, it in lst:
                k = B.kind_of(it)
                if k == "mc":
                    exp_ans.append("%s. %s" % (it["a"], it["o"][B.LETTERS.index(it["a"])]))
                elif k == "sa":
                    exp_ans.append(it["ans"])
        if len(got) != len(exp_ans):
            BAD.append("%s: số dòng “Đáp án:” là %d, cần %d" % (label, len(got), len(exp_ans)))
        else:
            for i, (g, e) in enumerate(zip(got, exp_ans)):
                if g.rstrip(".") != e.rstrip("."):
                    BAD.append("%s: dòng đáp án thứ %d lệch: “%s” ≠ “%s”" % (label, i + 1, g, e))

        # hình vẽ
        need = {it["fig"] for _c, _t, _n, lst in struct for _i, it in lst if it.get("fig")}
        nq, ns = len(images_in(qp)), len(images_in(sp))
        if nq < len(need) or ns < len(need):
            BAD.append("%s: số hình nhúng (%d/%d) ít hơn số hình cần (%d)"
                       % (label, nq, ns, len(need)))

        # bảng số liệu
        need_tbl = sum(1 for _c, _t, _n, lst in struct for _i, it in lst if it.get("tbl"))

        print("%s:" % label)
        print("   đề bài   : %d đoạn văn, %d câu, %d hình nhúng, %d bảng"
              % (len(tq), len(qids), nq, len(dq.tables)))
        print("   lời giải : %d đoạn văn, %d câu, %d hình nhúng, %d bảng"
              % (len(ts), len(sids), ns, len(ds.tables)))
        print("   hình khác nhau cần dùng: %d ; bảng số liệu trong đề: %d"
              % (len(need), need_tbl))

    if BAD:
        print("\n--- LỖI (%d) ---" % len(BAD))
        for b in BAD:
            print("  ✗", b)
        return 1
    print("\n✓ Bốn tệp .docx đồng bộ và đầy đủ.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
