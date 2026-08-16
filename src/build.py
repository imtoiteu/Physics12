# -*- coding: utf-8 -*-
"""Sinh hai file Word: đề cho học sinh và bản dành cho giáo viên."""

import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from part_a import PART_A
from part_b import PART_B
from part_c import PART_C
from part_d import PART_D

OUT_DIR = "/root/imtoiteu/Physics12"

SECTIONS = [
    ("PHẦN A – MỨC ĐỘ NHẬN BIẾT",
     "Kiểm tra việc ghi nhớ khái niệm, định nghĩa, đơn vị, hệ thức cơ bản.", PART_A),
    ("PHẦN B – MỨC ĐỘ THÔNG HIỂU",
     "Kiểm tra khả năng giải thích hiện tượng, đọc đồ thị và phân biệt các khái niệm gần nhau.", PART_B),
    ("PHẦN C – MỨC ĐỘ VẬN DỤNG",
     "Kiểm tra khả năng lập luận vật lí, xử lí tình huống thực tế, thí nghiệm và đồ thị.", PART_C),
    ("PHẦN D – MỨC ĐỘ VẬN DỤNG CAO (CÂU HỎI PHÂN LOẠI)",
     "Câu hỏi khó, nhiều phương án gây nhiễu dựa trên sai lầm phổ biến; dùng để phân loại học sinh giỏi – xuất sắc.", PART_D),
]

LETTERS = ["A", "B", "C", "D"]

SCOPE_NOTE = [
    ("Chương trình", "Vật lí 12 – Chương trình GDPT 2018 (áp dụng cho năm học 2025 – 2026)."),
    ("Phạm vi", "Học kì I: Chủ đề 1 – Vật lí nhiệt; Chủ đề 2 – Khí lí tưởng."),
    ("Nội dung Chủ đề 1",
     "Cấu trúc của chất rắn, chất lỏng, chất khí; sự chuyển thể (nóng chảy, hoá hơi, "
     "bay hơi, sôi, ngưng tụ); nội năng và định luật I nhiệt động lực học; thang nhiệt độ "
     "Celsius và Kelvin, nhiệt độ không tuyệt đối; nhiệt dung riêng, nhiệt nóng chảy riêng, "
     "nhiệt hoá hơi riêng; các thí nghiệm đo nhiệt dung riêng, nhiệt nóng chảy riêng, nhiệt hoá hơi riêng."),
    ("Nội dung Chủ đề 2",
     "Mô hình động học phân tử chất khí, chuyển động Brown, khí lí tưởng; định luật Boyle "
     "(quá trình đẳng nhiệt), định luật Charles (quá trình đẳng áp), quá trình đẳng tích; "
     "phương trình trạng thái của khí lí tưởng và phương trình Clapeyron – Mendeleev; "
     "áp suất chất khí và động năng phân tử theo mô hình động học phân tử; hằng số Boltzmann."),
    ("Mục tiêu", "Ôn luyện lí thuyết cho Kỳ thi tốt nghiệp THPT và các kỳ thi đánh giá năng lực."),
    ("Cấu trúc", "120 câu trắc nghiệm 4 phương án A, B, C, D; mỗi câu có duy nhất một đáp án đúng; "
                 "sắp xếp theo bốn mức độ từ dễ đến rất khó."),
]


def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = style.paragraph_format
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.0)


def para(doc, text="", bold=False, italic=False, size=12, align=None,
         space_before=0, space_after=2, indent=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"
        if color is not None:
            r.font.color.rgb = color
    return p


def title_block(doc, subtitle):
    para(doc, "NGÂN HÀNG CÂU HỎI LÍ THUYẾT", bold=True, size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "VẬT LÍ 12 – HỌC KÌ I", bold=True, size=16,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "Chương trình GDPT 2018 – Năm học 2025 – 2026", bold=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    para(doc, "Chủ đề: Vật lí nhiệt – Khí lí tưởng", italic=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, subtitle, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=6, color=RGBColor(0x00, 0x00, 0x00))
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("─" * 68).font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def scope_block(doc):
    para(doc, "I. PHẠM VI VÀ MỤC TIÊU CỦA NGÂN HÀNG CÂU HỎI", bold=True, size=13,
         space_before=4, space_after=4)
    for k, v in SCOPE_NOTE:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run("• " + k + ": ")
        r.bold = True
        r.font.size = Pt(12)
        r2 = p.add_run(v)
        r2.font.size = Pt(12)


def question_block(doc, idx, item, with_answer):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("Câu %d. " % idx)
    r.bold = True
    r.font.size = Pt(12)
    r2 = p.add_run(item["q"])
    r2.font.size = Pt(12)

    for i, opt in enumerate(item["o"]):
        op = doc.add_paragraph()
        op.paragraph_format.space_after = Pt(0)
        op.paragraph_format.left_indent = Cm(0.6)
        rr = op.add_run("%s. " % LETTERS[i])
        rr.bold = True
        rr.font.size = Pt(12)
        rr2 = op.add_run(opt)
        rr2.font.size = Pt(12)

    if with_answer:
        ans_i = LETTERS.index(item["a"])
        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(2)
        ap.paragraph_format.space_after = Pt(0)
        ap.paragraph_format.left_indent = Cm(0.6)
        ra = ap.add_run("Đáp án: %s. " % item["a"])
        ra.bold = True
        ra.font.size = Pt(12)
        ra.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        ra2 = ap.add_run(item["o"][ans_i])
        ra2.bold = True
        ra2.font.size = Pt(12)
        ra2.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        ep = doc.add_paragraph()
        ep.paragraph_format.space_before = Pt(1)
        ep.paragraph_format.space_after = Pt(2)
        ep.paragraph_format.left_indent = Cm(0.6)
        re = ep.add_run("Giải thích: ")
        re.bold = True
        re.italic = True
        re.font.size = Pt(12)
        re2 = ep.add_run(item["e"])
        re2.italic = True
        re2.font.size = Pt(12)


def answer_key_table(doc, all_items):
    para(doc, "BẢNG ĐÁP ÁN TỔNG HỢP", bold=True, size=13,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=6)
    ncols = 10
    nrows = (len(all_items) + ncols - 1) // ncols
    table = doc.add_table(rows=nrows * 2, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in range(nrows):
        for c in range(ncols):
            k = r * ncols + c
            head = table.cell(2 * r, c)
            body = table.cell(2 * r + 1, c)
            for cell in (head, body):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            if k < len(all_items):
                rh = head.paragraphs[0].add_run(str(k + 1))
                rh.font.size = Pt(10)
                rb = body.paragraphs[0].add_run(all_items[k]["a"])
                rb.bold = True
                rb.font.size = Pt(11)


def build(with_answer, filename, subtitle, note):
    doc = Document()
    set_base_style(doc)
    title_block(doc, subtitle)
    scope_block(doc)
    para(doc, note, italic=True, size=11.5, space_before=4, space_after=6)

    idx = 0
    all_items = []
    for sec_title, sec_desc, items in SECTIONS:
        para(doc, sec_title, bold=True, size=13.5, space_before=12, space_after=1,
             align=WD_ALIGN_PARAGRAPH.CENTER)
        para(doc, "(%d câu – %s)" % (len(items), sec_desc), italic=True, size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        for item in items:
            idx += 1
            all_items.append(item)
            question_block(doc, idx, item, with_answer)

    if with_answer:
        doc.add_page_break()
        answer_key_table(doc, all_items)

    para(doc, "--- HẾT ---", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=12)
    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    return path, idx


# Các câu có 4 phương án là số và được sắp xếp đơn điệu (tăng/giảm dần):
# giữ nguyên thứ tự phương án để bảng đáp án không bị lộn xộn về mặt trình bày.
KEEP_ORDER = {31, 37, 65, 82}


def rebalance():
    """Phân bố lại vị trí đáp án đúng cho đều giữa A, B, C, D.

    Chỉ hoán vị vị trí của các phương án trong cùng một câu (nội dung không đổi),
    thực hiện một lần trên dữ liệu dùng chung nên hai file luôn đồng bộ.
    """
    items = []
    for _, _, sec in SECTIONS:
        items.extend(sec)
    counts = {L: 0 for L in LETTERS}
    for it in items:
        counts[it["a"]] += 1
    target = len(items) // 4

    order = [(k * 7) % len(items) for k in range(len(items))]
    for i in order:
        if i + 1 in KEEP_ORDER:
            continue
        cur = items[i]["a"]
        if counts[cur] <= target:
            continue
        need = [L for L in LETTERS if counts[L] < target]
        if not need:
            break
        dst = need[0]
        a, b = LETTERS.index(cur), LETTERS.index(dst)
        opts = items[i]["o"]
        opts[a], opts[b] = opts[b], opts[a]
        items[i]["a"] = dst
        counts[cur] -= 1
        counts[dst] += 1
    return counts


def validate():
    seen = {}
    problems = []
    counts = {"A": 0, "B": 0, "C": 0, "D": 0}
    n = 0
    for sec_title, _, items in SECTIONS:
        for it in items:
            n += 1
            if len(it["o"]) != 4:
                problems.append("Câu %d: không có đúng 4 phương án" % n)
            if it["a"] not in LETTERS:
                problems.append("Câu %d: đáp án không hợp lệ" % n)
            counts[it["a"]] += 1
            if len(set(o.strip().lower() for o in it["o"])) != 4:
                problems.append("Câu %d: có phương án trùng nhau" % n)
            key = it["q"].strip().lower()[:60]
            if key in seen:
                problems.append("Câu %d trùng phần dẫn với câu %d" % (n, seen[key]))
            seen[key] = n
            if not it.get("e") or len(it["e"]) < 40:
                problems.append("Câu %d: giải thích quá ngắn" % n)
    return n, counts, problems


if __name__ == "__main__":
    before = {L: 0 for L in LETTERS}
    for _, _, sec in SECTIONS:
        for it in sec:
            before[it["a"]] += 1
    print("Phân bố đáp án trước khi cân bằng:", before)
    print("Phân bố đáp án sau khi cân bằng:  ", rebalance())

    n, counts, problems = validate()
    print("Tổng số câu:", n)
    print("Phân bố đáp án:", counts)
    print("Lỗi phát hiện:", problems if problems else "không có")

    p1, c1 = build(
        False,
        "NGAN_HANG_CAU_HOI_VAT_LI_12_HK1_DE_HOC_SINH.docx",
        "BẢN DÀNH CHO HỌC SINH",
        "Hướng dẫn: Mỗi câu hỏi chỉ có một phương án đúng. Hãy chọn phương án mà em cho là đúng nhất. "
        "Các câu được sắp xếp theo mức độ khó tăng dần; nên làm tuần tự từ Phần A đến Phần D."
    )
    p2, c2 = build(
        True,
        "NGAN_HANG_CAU_HOI_VAT_LI_12_HK1_BAN_GIAO_VIEN.docx",
        "BẢN DÀNH CHO GIÁO VIÊN (CÓ ĐÁP ÁN VÀ GIẢI THÍCH)",
        "Lưu ý sử dụng: Bản này giữ nguyên thứ tự và nội dung câu hỏi như bản dành cho học sinh. "
        "Mỗi câu kèm đáp án đúng và phần giải thích cơ sở vật lí, đồng thời chỉ rõ sai lầm mà các "
        "phương án nhiễu nhắm tới. Bảng đáp án tổng hợp được đặt ở cuối tài liệu."
    )
    print("Đã tạo:", p1, "(%d câu)" % c1)
    print("Đã tạo:", p2, "(%d câu)" % c2)
