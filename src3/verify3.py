# -*- coding: utf-8 -*-
"""Kiểm tra chất lượng bộ tài liệu học kì I trước khi bàn giao."""

import os, re, sys, difflib, collections

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
FIGS = os.path.join(ROOT, "figs")
sys.path.insert(0, HERE)

import build3
from build3 import (LETTERS, kind_of, explanation, theory_structure, calc_structure,
                    rebalance_answers, short_answer)
from ch1_theory import CH1
from ch2_theory import CH2
from ch1_concept import MC1, DS1
from ch2_concept import MC2, DS2
from ch1_calc import CALC1
from ch2_calc import CALC2
from combined_calc12 import CALC12

errors, warns = [], []


def err(msg):
    errors.append(msg)


def warn(msg):
    warns.append(msg)


# --------------------------------------------------------------- 1. cấu trúc
rebalance_answers(MC1, MC2, CALC1, CALC2, CALC12)
ths, cas = theory_structure(), calc_structure()
allitems = [(i, it) for _t, secs in ths + cas for _s, lst in secs for i, it in lst]

kinds = collections.Counter(kind_of(it) for _i, it in allitems)
dist = collections.Counter(it["a"] for _i, it in allitems if kind_of(it) == "mc")

for ident, it in allitems:
    k = kind_of(it)
    if k == "mc":
        if len(it["o"]) != 4:
            err("%s: có %d phương án (phải là 4)" % (ident, len(it["o"])))
        if it["a"] not in LETTERS:
            err("%s: đáp án '%s' không hợp lệ" % (ident, it["a"]))
        if len(set(o.strip().lower() for o in it["o"])) != 4:
            err("%s: có hai phương án trùng nội dung" % ident)
        for o in it["o"]:
            if not o.strip():
                err("%s: có phương án rỗng" % ident)
    elif k == "ds":
        if len(it["items"]) != 4:
            err("%s: có %d ý (phải là 4)" % (ident, len(it["items"])))
        if all(v for _t, v, _e in it["items"]) or not any(v for _t, v, _e in it["items"]):
            warn("%s: cả 4 ý cùng đúng hoặc cùng sai" % ident)
        for txt, _v, ex in it["items"]:
            if len(ex.strip()) < 25:
                warn("%s: giải thích quá ngắn cho ý '%s...'" % (ident, txt[:35]))
    else:
        if not str(it.get("ans", "")).strip():
            err("%s: thiếu đáp án" % ident)
    if k != "ds" and len(explanation(it).strip()) < 25:
        err("%s: thiếu lời giải hoặc lời giải quá ngắn" % ident)

# --------------------------------------------------- 2. dẫn chiếu chữ cái
LETTER_REF = re.compile(
    r"(?:[Pp]hương án|[Đđ]áp án|[Cc]âu|[Ýý])\s+[ABCD]\b"
    r"|\b[ABCD]\s*(?:,\s*[ABCD]\s*)*(?:và|hoặc)\s+[ABCD]\b"
    r"|[Cc]họn\s+[ABCD]\b")
for ident, it in allitems:
    texts = [explanation(it)]
    if kind_of(it) == "ds":
        texts += [e for _t, _v, e in it["items"]]
    for t in texts:
        m = LETTER_REF.search(t)
        if m:
            err("%s: lời giải dẫn chiếu chữ cái phương án -> '%s'" % (ident, m.group(0)))

# --------------------------------------------------- 3. trùng lặp phần dẫn
def stem_of(it):
    return (it["stem"] if kind_of(it) == "ds" else it["q"]).strip()


norm = lambda s: re.sub(r"[^0-9a-zà-ỹ]+", " ", s.lower()).strip()
stems = [(i, norm(stem_of(it))) for i, it in allitems]
for a in range(len(stems)):
    for b in range(a + 1, len(stems)):
        r = difflib.SequenceMatcher(None, stems[a][1], stems[b][1]).ratio()
        if r > 0.90:
            err("Phần dẫn gần trùng nhau (%.2f): %s  ↔  %s" % (r, stems[a][0], stems[b][0]))

# --------------------------------------------------- 4. hình vẽ
used = set()
for blocks in (CH1, CH2):
    for blk in blocks:
        if blk[0] == "fig":
            used.add(blk[1])
for _i, it in allitems:
    if it.get("fig"):
        used.add(it["fig"])
for name in sorted(used):
    if not os.path.exists(os.path.join(FIGS, name + ".png")):
        err("Thiếu file hình: %s.png" % name)

import figures1
made = set()
for f in figures1.ALL:
    made.add(f.__name__.split("_", 1)[1])
declared = {n for n in os.listdir(FIGS) if n.startswith("h") and n[1:3].isdigit()}
declared = {n[:-4] for n in declared}
unused = declared - used
if unused:
    warn("Hình đã vẽ nhưng chưa dùng ở đâu: %s" % ", ".join(sorted(unused)))

# --------------------------------------------------- 5. đồng bộ với file .docx
from docx import Document

def idents_in(path, pattern):
    doc = Document(os.path.join(ROOT, path))
    out = []
    for p in doc.paragraphs:
        t = p.text.strip()
        m = pattern.match(t)
        if m:
            out.append(m.group(1))
    return out


PAT_Q = re.compile(r"^((?:Câu ĐS|Câu|Bài) [\d.]+)\.")
f2 = idents_in("HK1_FILE2_BAI_TAP_LI_THUYET_CHUONG_1_VA_2.docx", PAT_Q)
f3 = idents_in("HK1_FILE3_BAI_TAP_TINH_TOAN_VA_SUY_LUAN.docx", PAT_Q)
f4 = idents_in("HK1_FILE4_DAP_AN_VA_LOI_GIAI_CHI_TIET.docx", PAT_Q)
expected = [i for i, _ in allitems]
if f2 + f3 != f4:
    err("Mã câu trong file đề (%d) KHÔNG khớp file lời giải (%d)" % (len(f2 + f3), len(f4)))
if f2 + f3 != expected:
    err("Mã câu trong file đề không khớp cấu trúc dựng (%d ↔ %d)"
        % (len(f2 + f3), len(expected)))

# file học sinh không được lộ đáp án
doc2 = Document(os.path.join(ROOT, "HK1_FILE2_BAI_TAP_LI_THUYET_CHUONG_1_VA_2.docx"))
doc3 = Document(os.path.join(ROOT, "HK1_FILE3_BAI_TAP_TINH_TOAN_VA_SUY_LUAN.docx"))
for nm, d in [("FILE2", doc2), ("FILE3", doc3)]:
    txt = "\n".join(p.text for p in d.paragraphs)
    for bad in ("Đáp án:", "Lời giải:", "Giải thích:", "→ ĐÚNG", "→ SAI"):
        if bad in txt:
            err("%s (file đề) bị lộ nội dung '%s'" % (nm, bad))

# --------------------------------------------------- 6. hình nhúng trong docx
def n_images(path):
    import zipfile
    with zipfile.ZipFile(os.path.join(ROOT, path)) as z:
        return len([n for n in z.namelist() if n.startswith("word/media/")])


files = ["HK1_FILE1_TAI_LIEU_DAY_HOC_CHUONG_1_VA_2.docx",
         "HK1_FILE2_BAI_TAP_LI_THUYET_CHUONG_1_VA_2.docx",
         "HK1_FILE3_BAI_TAP_TINH_TOAN_VA_SUY_LUAN.docx",
         "HK1_FILE4_DAP_AN_VA_LOI_GIAI_CHI_TIET.docx"]

# --------------------------------------------------- BÁO CÁO
print("Loại câu:", dict(kinds), " | tổng:", sum(kinds.values()))
print("Phân bố đáp án:", dict(dist))
print("Số hình gốc đã vẽ:", len(figures1.ALL), "| số hình được dùng:", len(used))
print("Đồng bộ đề ↔ lời giải:", f2 + f3 == f4 == expected)
for f in files:
    p = os.path.join(ROOT, f)
    print("   %-52s %3d hình  %6d KB" % (f, n_images(f), os.path.getsize(p) // 1024))

print("\nCẢNH BÁO:", len(warns))
for w in warns:
    print("   ⚠", w)
print("\nLỖI:", len(errors))
for e in errors:
    print("   ✗", e)
if not errors:
    print("   KHÔNG CÓ LỖI CẤU TRÚC")
