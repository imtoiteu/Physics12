# -*- coding: utf-8 -*-
"""Sinh 4 tài liệu Word cho Vật lí 12 học kì I (Chương I và Chương II)."""

import os, re, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.abspath(os.path.join(HERE, ".."))
FIGS = os.path.join(ROOT, "figs")
sys.path.insert(0, HERE)

from ch1_theory import CH1
from ch2_theory import CH2
from ch1_concept import MC1, DS1
from ch2_concept import MC2, DS2
from ch1_calc import CALC1
from ch2_calc import CALC2
from combined_calc12 import CALC12

LETTERS = ["A", "B", "C", "D"]
DS_LABELS = ["a)", "b)", "c)", "d)"]
FONT = "Times New Roman"
DARKRED = RGBColor(0xA6, 0x1B, 0x1B)
DARKBLUE = RGBColor(0x1F, 0x3A, 0x6E)
DARKGREEN = RGBColor(0x1E, 0x64, 0x36)


# ----------------------------------------------------------------- tiện ích
def set_base(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(3)
    st.paragraph_format.line_spacing = 1.13
    for s in doc.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2); s.right_margin = Cm(1.8)


def P(doc, text="", bold=False, italic=False, size=12, align=None, before=0, after=3,
      indent=None, color=None, first_line=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if indent is not None:
        pf.left_indent = Cm(indent)
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)
    if text:
        for i, line in enumerate(str(text).split("\n")):
            if i:
                p.add_run().add_break()
            r = p.add_run(line)
            r.bold = bold; r.italic = italic
            r.font.size = Pt(size); r.font.name = FONT
            if color is not None:
                r.font.color.rgb = color
    return p


def rich(doc, parts, indent=None, before=0, after=3, size=12):
    """parts = [(text, bold, italic, color|None), ...]"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if indent is not None:
        pf.left_indent = Cm(indent)
    for text, b, i, c in parts:
        for k, line in enumerate(str(text).split("\n")):
            if k:
                p.add_run().add_break()
            r = p.add_run(line)
            r.bold = b; r.italic = i; r.font.size = Pt(size); r.font.name = FONT
            if c is not None:
                r.font.color.rgb = c
    return p


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def box(doc, title, text, fill="EEF3FA", tcolor=DARKBLUE):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, fill)
    c.paragraphs[0].paragraph_format.space_after = Pt(2)
    if title:
        r = c.paragraphs[0].add_run(title)
        r.bold = True; r.font.size = Pt(11.5); r.font.name = FONT; r.font.color.rgb = tcolor
    tp = c.add_paragraph()
    tp.paragraph_format.space_after = Pt(1)
    for i, line in enumerate(str(text).split("\n")):
        if i:
            tp.add_run().add_break()
        r = tp.add_run(line)
        r.font.size = Pt(11.5); r.font.name = FONT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def picture(doc, name, caption=None, maxw=13.0, maxh=8.6):
    path = os.path.join(FIGS, name + ".png")
    if not os.path.exists(path):
        raise FileNotFoundError(path)
    with Image.open(path) as im:
        w, h = im.size
    width = maxw
    if h / w * width > maxh:
        width = maxh * w / h
    doc.add_picture(path, width=Cm(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(4)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
    if caption:
        P(doc, caption, italic=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)


def table(doc, caption, headers, rows):
    if caption:
        P(doc, caption, italic=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER,
          before=4, after=2)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        c = t.rows[0].cells[j]
        shade(c, "DCE6F1")
        c.paragraphs[0].paragraph_format.space_after = Pt(1)
        r = c.paragraphs[0].add_run(htxt)
        r.bold = True; r.font.size = Pt(10.5); r.font.name = FONT
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].paragraphs[0].paragraph_format.space_after = Pt(1)
            r = cells[j].paragraphs[0].add_run(str(val))
            r.font.size = Pt(10.5); r.font.name = FONT
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def cover(doc, title_lines, subtitle, note):
    P(doc, "TÀI LIỆU DẠY HỌC VÀ ÔN THI", bold=True, size=12.5,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    P(doc, "VẬT LÍ 12 – HỌC KÌ I", bold=True, size=19,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    P(doc, "Chương trình GDPT 2018 – Năm học 2026 – 2027", bold=True, size=12,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    P(doc, "CHƯƠNG I: VẬT LÍ NHIỆT  •  CHƯƠNG II: KHÍ LÍ TƯỞNG", italic=True, size=11.5,
      align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    for i, line in enumerate(title_lines):
        P(doc, line, bold=True, size=15 if i == 0 else 13,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=1, color=DARKRED)
    P(doc, subtitle, italic=True, size=11.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    box(doc, "ĐỊNH HƯỚNG BIÊN SOẠN", note, fill="F4F6F8")


ORIENT = ("Tài liệu bám sát yêu cầu cần đạt của Chương trình GDPT 2018 và định dạng đề thi "
          "tốt nghiệp THPT từ năm 2025: Phần I gồm câu trắc nghiệm nhiều phương án lựa chọn, "
          "Phần II gồm câu trắc nghiệm đúng/sai (mỗi câu 4 ý), Phần III gồm câu trả lời ngắn.\n"
          "Trọng tâm đánh giá: hiểu bản chất hiện tượng nhiệt, điều kiện áp dụng của các định "
          "luật chất khí, đọc hiểu đồ thị – bảng số liệu, suy luận thực nghiệm và vận dụng vào "
          "bối cảnh thực tiễn.")


# ----------------------------------------------------- đánh số thống nhất
def theory_structure():
    """[(tên phần, [(tên mục, [(mã câu, item)])])] cho tài liệu bài tập lí thuyết."""
    out = []
    for mc, ds, tag, title in [
            (MC1, DS1, "1", "PHẦN A – CHƯƠNG I: VẬT LÍ NHIỆT"),
            (MC2, DS2, "2", "PHẦN B – CHƯƠNG II: KHÍ LÍ TƯỞNG")]:
        secs = []
        n = 0
        for level, items in mc.items():
            lst = []
            for it in items:
                n += 1
                lst.append(("Câu %s.%d" % (tag, n), it))
            secs.append((level, lst))
        lst = []
        for k, it in enumerate(ds, 1):
            lst.append(("Câu ĐS %s.%d" % (tag, k), it))
        secs.append(("Câu trắc nghiệm ĐÚNG/SAI (mỗi câu gồm 4 ý)", lst))
        out.append((title, secs))
    return out


def calc_structure():
    out = []
    for data, tag, title in [
            (CALC1, "1", "PHẦN A – BÀI TẬP CHƯƠNG I: VẬT LÍ NHIỆT"),
            (CALC2, "2", "PHẦN B – BÀI TẬP CHƯƠNG II: KHÍ LÍ TƯỞNG"),
            (CALC12, "12", "PHẦN C – BÀI TẬP TỔNG HỢP CHƯƠNG I + CHƯƠNG II")]:
        secs = []
        n = 0
        for level, items in data.items():
            lst = []
            for it in items:
                n += 1
                lst.append(("Bài %s.%d" % (tag, n), it))
            secs.append((level, lst))
        out.append((title, secs))
    return out


SUP = {"⁰": "0", "¹": "1", "²": "2", "³": "3", "⁴": "4",
       "⁵": "5", "⁶": "6", "⁷": "7", "⁸": "8", "⁹": "9", "⁻": "-"}

_NUM = re.compile(r"^(?P<sig>[-−]?\d+(?:[.,]\d+)?)(?:\s*[·x×]\s*10(?P<exp>[⁰¹²³⁴⁵⁶⁷⁸⁹⁻]+))?")


def parse_value(opt):
    """Đọc giá trị số ở đầu một phương án. Trả về None nếu không chắc chắn.

    Chấp nhận các dạng: '420 kJ.', '3,0·10⁵ Pa.', '≈ 6,21·10⁻²¹ J.', '−250 J.'
    Từ chối các dạng mập mờ như '1/√10 ≈ 0,316.' để không sắp xếp nhầm."""
    t = opt.strip().lstrip("≈~±").strip()
    m = _NUM.match(t)
    if not m:
        return None
    rest = t[m.end():]
    # phần còn lại phải là đơn vị / dấu chấm câu, không được là biểu thức tiếp diễn
    if rest[:1] in {"/", "√", "^", ":", "-", "−", "·", "+"}:
        return None
    try:
        val = float(m.group("sig").replace("−", "-").replace(",", "."))
    except ValueError:
        return None
    if m.group("exp"):
        e = "".join(SUP.get(ch, "") for ch in m.group("exp"))
        try:
            val *= 10.0 ** int(e)
        except ValueError:
            return None
    return val


def sort_numeric_options(item):
    """Sắp xếp các phương án đáp số theo thứ tự tăng dần và cập nhật đáp án đúng.

    Vừa giúp học sinh dễ đối chiếu, vừa làm vị trí đáp án đúng phân tán tự nhiên
    thay vì luôn nằm ở một chữ cái. Trả về True nếu đã sắp xếp."""
    opts = item["o"]
    vals = [parse_value(o) for o in opts]
    if any(v is None for v in vals) or len(set(vals)) != len(vals):
        return False
    key_text = opts[LETTERS.index(item["a"])]
    order = sorted(range(len(opts)), key=lambda i: vals[i])
    item["o"] = [opts[i] for i in order]
    item["a"] = LETTERS[item["o"].index(key_text)]
    return True


def rebalance_answers(*datasets):
    """Phân bố lại vị trí đáp án đúng cho đều giữa A, B, C, D.

    Chỉ hoán vị vị trí các phương án trong cùng một câu (nội dung không đổi) và chỉ
    áp dụng cho những câu có phương án dạng chữ, để các câu đáp số giữ nguyên thứ tự
    vốn dễ đọc hơn."""
    pool, counts, nsorted = [], {L: 0 for L in LETTERS}, 0
    for data in datasets:
        for _lvl, items in data.items():
            for it in items:
                if "o" not in it:
                    continue
                if sort_numeric_options(it):
                    nsorted += 1          # câu đáp số: xếp tăng dần, không hoán vị thêm
                else:
                    pool.append(it)       # câu phương án chữ: dùng để cân bằng
                counts[it["a"]] += 1
    print("  (đã sắp xếp tăng dần %d câu đáp số, còn %d câu chữ để cân bằng)"
          % (nsorted, len(pool)))
    target = sum(counts.values()) // 4
    order = [(k * 7) % len(pool) for k in range(len(pool))]
    for i in order:
        it = pool[i]
        cur = it["a"]
        if counts[cur] <= target:
            continue
        need = [L for L in LETTERS if counts[L] < target]
        if not need:
            break
        dst = need[0]
        a, b = LETTERS.index(cur), LETTERS.index(dst)
        it["o"][a], it["o"][b] = it["o"][b], it["o"][a]
        it["a"] = dst
        counts[cur] -= 1; counts[dst] += 1
    return counts


def kind_of(item):
    if "o" in item:
        return "mc"
    if "items" in item:
        return "ds"
    return "short"


def explanation(item):
    """Chấp nhận cả khoá 'sol' (file tính toán) lẫn 'e' (file lí thuyết)."""
    return item.get("sol") or item.get("e") or ""


# ----------------------------------------------------- hiển thị câu hỏi
def _stem_extras(doc, item, maxw=11.5, maxh=7.2):
    if item.get("fig"):
        picture(doc, item["fig"], maxw=maxw, maxh=maxh)
    if item.get("tbl"):
        cap, hdr, rows = item["tbl"]
        table(doc, cap, hdr, rows)


def render_question(doc, ident, item, show_answer=False):
    k = kind_of(item)
    if k == "ds":
        rich(doc, [(ident + ". ", True, False, None), (item["stem"], False, False, None)],
             before=7, after=2)
        _stem_extras(doc, item)
        for lab, (txt, val, _ex) in zip(DS_LABELS, item["items"]):
            parts = [(lab + " ", True, False, None), (txt, False, False, None)]
            if show_answer:
                parts.append(("   → " + ("ĐÚNG" if val else "SAI"), True, False,
                              DARKBLUE if val else DARKRED))
            rich(doc, parts, indent=0.55, after=1)
        return
    rich(doc, [(ident + ". ", True, False, None), (item["q"], False, False, None)],
         before=7, after=2)
    _stem_extras(doc, item)
    if k == "mc":
        for i, opt in enumerate(item["o"]):
            rich(doc, [("%s. " % LETTERS[i], True, False, None), (opt, False, False, None)],
                 indent=0.55, after=0)
    else:
        P(doc, "...................................................................................",
          size=11, indent=0.55, after=1)


def render_solution(doc, ident, item):
    k = kind_of(item)
    if k == "ds":
        rich(doc, [(ident + ". ", True, False, None), (item["stem"], False, True, None)],
             before=7, after=2)
        _stem_extras(doc, item, maxw=10.5, maxh=6.6)
        for lab, (txt, val, ex) in zip(DS_LABELS, item["items"]):
            rich(doc, [(lab + " ", True, False, None), (txt, False, False, None),
                       ("  →  " + ("ĐÚNG" if val else "SAI"), True, False,
                        DARKBLUE if val else DARKRED)], indent=0.55, after=1)
            rich(doc, [("Giải thích: ", True, True, None), (ex, False, True, None)],
                 indent=1.0, after=2, size=11.5)
        return

    rich(doc, [(ident + ". ", True, False, None), (item["q"], False, True, None)],
         before=7, after=2)
    _stem_extras(doc, item, maxw=10.5, maxh=6.6)
    if k == "mc":
        ans = item["a"]
        txt = item["o"][LETTERS.index(ans)]
        rich(doc, [("Đáp án: %s. " % ans, True, False, DARKRED), (txt, True, False, DARKRED)],
             indent=0.55, after=2)
    else:
        rich(doc, [("Đáp án: ", True, False, DARKRED), (item["ans"], True, False, DARKRED)],
             indent=0.55, after=2)
    rich(doc, [("Lời giải: ", True, True, None), (explanation(item), False, False, None)],
         indent=0.55, after=2)


# ----------------------------------------------------- FILE 1: lí thuyết
def build_teaching():
    doc = Document(); set_base(doc)
    cover(doc,
          ["FILE 1 – TÀI LIỆU DẠY HỌC", "LÍ THUYẾT CHƯƠNG I VÀ CHƯƠNG II"],
          "Dùng cho giáo viên giảng dạy trên lớp và học sinh hệ thống hoá kiến thức",
          ORIENT)
    doc.add_page_break()

    for blocks in (CH1, CH2):
        for blk in blocks:
            t = blk[0]
            if t == "h1":
                P(doc, blk[1], bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
                  before=6, after=8, color=DARKRED)
            elif t == "h2":
                P(doc, blk[1], bold=True, size=13.5, before=12, after=4, color=DARKBLUE)
            elif t == "h3":
                P(doc, blk[1], bold=True, size=12.5, before=8, after=3)
            elif t == "p":
                P(doc, blk[1], after=4, first_line=0.6)
            elif t == "b":
                rich(doc, [("• ", True, False, None), (blk[1], False, False, None)],
                     indent=0.5, after=2)
            elif t == "f":
                P(doc, blk[1], bold=True, size=12.5, align=WD_ALIGN_PARAGRAPH.CENTER,
                  before=4, after=4, color=DARKRED)
            elif t == "box":
                box(doc, blk[1], blk[2])
            elif t == "trap":
                box(doc, "⚠ SAI LẦM THƯỜNG GẶP – PHÂN BIỆT CHO RÕ", blk[1],
                    fill="FDF0EF", tcolor=DARKRED)
            elif t == "exam":
                box(doc, "★ LƯU Ý KHI ÔN THI", blk[1], fill="F2F8F0", tcolor=DARKGREEN)
            elif t == "fig":
                picture(doc, blk[1], blk[2])
            elif t == "tbl":
                table(doc, blk[1], blk[2], blk[3])
        doc.add_page_break()

    P(doc, "--- HẾT FILE 1 ---", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=10)
    path = os.path.join(ROOT, "HK1_FILE1_TAI_LIEU_DAY_HOC_CHUONG_1_VA_2.docx")
    doc.save(path)
    return path


# ----------------------------------------------------- FILE 2 & 3: đề bài
def build_exercises(structure, filename, title_lines, subtitle, extra_note=None):
    doc = Document(); set_base(doc)
    cover(doc, title_lines, subtitle, ORIENT)
    if extra_note:
        box(doc, "HƯỚNG DẪN SỬ DỤNG", extra_note, fill="F2F8F0", tcolor=DARKGREEN)
    doc.add_page_break()
    total = 0
    for part_title, secs in structure:
        P(doc, part_title, bold=True, size=14.5, align=WD_ALIGN_PARAGRAPH.CENTER,
          before=8, after=6, color=DARKRED)
        for sec_title, lst in secs:
            P(doc, "%s  (%d câu)" % (sec_title, len(lst)), bold=True, size=12.5,
              before=10, after=3, color=DARKBLUE)
            for ident, item in lst:
                render_question(doc, ident, item)
                total += 1
        doc.add_page_break()
    P(doc, "--- HẾT ---", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=8)
    path = os.path.join(ROOT, filename)
    doc.save(path)
    return path, total


# ----------------------------------------------------- FILE 4: lời giải
def answer_key(doc, title, pairs):
    P(doc, title, bold=True, size=12.5, before=10, after=4, color=DARKBLUE)
    ncol = 6
    nrow = (len(pairs) + ncol - 1) // ncol
    t = doc.add_table(rows=nrow, cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (ident, ans) in enumerate(pairs):
        c = t.rows[i // ncol].cells[i % ncol]
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.paragraphs[0].add_run(ident.replace("Câu ", "").replace("Bài ", ""))
        r.font.size = Pt(9.5); r.font.name = FONT
        c.paragraphs[0].add_run(": ").font.size = Pt(9.5)
        r2 = c.paragraphs[0].add_run(ans)
        r2.bold = True; r2.font.size = Pt(9.5); r2.font.name = FONT
        r2.font.color.rgb = DARKRED
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def short_answer(item):
    k = kind_of(item)
    if k == "mc":
        return item["a"]
    if k == "ds":
        return "".join("Đ" if v else "S" for _t, v, _e in item["items"])
    return item["ans"]


def build_solutions(theory_struct, calc_struct):
    doc = Document(); set_base(doc)
    cover(doc,
          ["FILE 4 – ĐÁP ÁN VÀ LỜI GIẢI CHI TIẾT", "cho FILE 2 và FILE 3"],
          "Mã số câu hỏi được giữ nguyên hoàn toàn so với hai file đề bài",
          ORIENT)
    box(doc, "CÁCH ĐỌC TÀI LIỆU NÀY",
        "• Mỗi câu được ghi lại nguyên văn phần dẫn (in nghiêng) kèm đáp án và lời giải.\n"
        "• Với câu đúng/sai, đáp án được ghi theo thứ tự bốn ý a) b) c) d), ví dụ “ĐĐSĐ”.\n"
        "• Bảng đáp án nhanh của từng phần được đặt ngay đầu phần tương ứng.\n"
        "• Với bài tự luận nhiều ý, dòng “Đáp án” chỉ ghi kết quả chính; phần “Lời giải” "
        "trình bày đầy đủ từng ý a), b), c), d).",
        fill="F4F6F8")
    doc.add_page_break()

    for big_title, struct in [("PHẦN I – LỜI GIẢI BÀI TẬP LÍ THUYẾT (FILE 2)", theory_struct),
                              ("PHẦN II – LỜI GIẢI BÀI TẬP TÍNH TOÁN (FILE 3)", calc_struct)]:
        P(doc, big_title, bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER,
          before=8, after=8, color=DARKRED)
        for part_title, secs in struct:
            P(doc, part_title, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
              before=10, after=5, color=DARKRED)
            pairs = [(ident, short_answer(item)) for _s, lst in secs for ident, item in lst]
            answer_key(doc, "BẢNG ĐÁP ÁN NHANH", pairs)
            for sec_title, lst in secs:
                P(doc, sec_title, bold=True, size=12.5, before=10, after=3, color=DARKBLUE)
                for ident, item in lst:
                    render_solution(doc, ident, item)
            doc.add_page_break()

    P(doc, "--- HẾT FILE 4 ---", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=8)
    path = os.path.join(ROOT, "HK1_FILE4_DAP_AN_VA_LOI_GIAI_CHI_TIET.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    before = {L: 0 for L in LETTERS}
    for data in (MC1, MC2, CALC1, CALC2, CALC12):
        for _l, items in data.items():
            for it in items:
                if "o" in it:
                    before[it["a"]] += 1
    print("Phân bố đáp án trước khi cân bằng:", before)
    print("Phân bố đáp án sau khi cân bằng:  ",
          rebalance_answers(MC1, MC2, CALC1, CALC2, CALC12))

    ths = theory_structure()
    cas = calc_structure()

    p1 = build_teaching()
    print("Đã tạo:", os.path.basename(p1))

    p2, n2 = build_exercises(
        ths, "HK1_FILE2_BAI_TAP_LI_THUYET_CHUONG_1_VA_2.docx",
        ["FILE 2 – BÀI TẬP LÍ THUYẾT", "CHƯƠNG I VÀ CHƯƠNG II"],
        "Câu hỏi khái niệm, hiện tượng, đồ thị và thí nghiệm – sắp xếp từ dễ đến rất khó",
        "Các câu được sắp xếp theo bốn mức độ tăng dần trong từng chương. "
        "Phần câu hỏi đúng/sai ở cuối mỗi chương mô phỏng đúng dạng câu hỏi Phần II của đề thi "
        "tốt nghiệp THPT: mỗi câu có bốn ý, thí sinh phải đánh giá đúng hoặc sai cho từng ý.")
    print("Đã tạo:", os.path.basename(p2), "-", n2, "câu")

    p3, n3 = build_exercises(
        cas, "HK1_FILE3_BAI_TAP_TINH_TOAN_VA_SUY_LUAN.docx",
        ["FILE 3 – BÀI TẬP TÍNH TOÁN VÀ SUY LUẬN", "CHƯƠNG I, CHƯƠNG II VÀ TỔNG HỢP"],
        "Bài tập định lượng theo ba dạng thức của đề thi, kèm bài tập tự luận vận dụng cao",
        "Dạng 1 tương ứng Phần I của đề thi (trắc nghiệm nhiều phương án lựa chọn); "
        "Dạng 2 tương ứng Phần II (đúng/sai); Dạng 3 tương ứng Phần III (trả lời ngắn). "
        "Dạng 4 là bài tập tự luận nhiều ý, dùng để rèn kĩ năng trình bày và tư duy tổng hợp, "
        "phù hợp cho ôn thi đánh giá năng lực và bồi dưỡng học sinh giỏi.")
    print("Đã tạo:", os.path.basename(p3), "-", n3, "bài")

    p4 = build_solutions(ths, cas)
    print("Đã tạo:", os.path.basename(p4))
    print("Tổng số câu hỏi/bài tập:", n2 + n3)
